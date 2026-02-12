"""
Document processing service - extracts text from PDF, DOCX, and TXT files
"""
import os
from typing import Optional, Dict
import PyPDF2
from docx import Document as DocxDocument


class DocumentProcessor:
    """Extract text content from various document formats"""

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> Optional[str]:
        """
        Extract text from a document file

        Args:
            file_path: Path to the document file
            file_type: Type of file ('pdf', 'docx', 'txt')

        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            if file_type == 'pdf':
                return DocumentProcessor._extract_from_pdf(file_path)
            elif file_type == 'docx':
                return DocumentProcessor._extract_from_docx(file_path)
            elif file_type == 'txt':
                return DocumentProcessor._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n\n'.join(text)

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return '\n\n'.join(paragraphs)

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    @staticmethod
    def get_metadata(file_path: str, file_type: str) -> Dict:
        """
        Extract metadata from document

        Args:
            file_path: Path to the document file
            file_type: Type of file ('pdf', 'docx', 'txt')

        Returns:
            Dictionary with metadata (page count, file size, etc.)
        """
        metadata = {
            'file_size': os.path.getsize(file_path)
        }

        try:
            if file_type == 'pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)
            elif file_type == 'docx':
                doc = DocxDocument(file_path)
                metadata['paragraph_count'] = len(doc.paragraphs)
        except Exception as e:
            print(f"Error extracting metadata: {str(e)}")

        return metadata
